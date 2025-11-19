# import streamlit as st
# import os
# import requests
# import json
# import numpy as np
# from sklearn.metrics.pairwise import cosine_similarity
# from io import BytesIO

# # For PDFs
# from langchain_community.document_loaders.pdf import PDFPlumberLoader
# # For DOCX
# import docx
# # For images (OCR)
# from PIL import Image
# import pytesseract

# # For text splitting
# from langchain_text_splitters import RecursiveCharacterTextSplitter

# # For exporting chat
# from fpdf import FPDF

# # ------------------------- API SETUP -------------------------
# API_KEY = os.getenv("GEMINI_API_KEY")   # MUST be set in Windows env
# if not API_KEY:
#     st.error("❌ GEMINI_API_KEY not set in environment variables.")
#     st.stop()

# EMBED_URL = f"https://generativelanguage.googleapis.com/v1beta/models/text-embedding-004:embedContent?key={API_KEY}"
# CHAT_URL  = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={API_KEY}"

# # ------------------------- TEMPLATE -------------------------
# template = """
# You are an assistant for question-answering using only the context below.
# If the answer is not in the context, say "I don't know."

# Question: {question}

# Context:
# {context}

# Answer:
# """

# # ------------------------- EMBEDDING FUNCTION -------------------------
# def get_embedding(text):
#     headers = {"Content-Type": "application/json"}
#     body = {"model": "models/text-embedding-004", "content": {"parts": [{"text": text}]}}
#     r = requests.post(EMBED_URL, headers=headers, json=body)
#     try:
#         return np.array(r.json()["embedding"]["values"])
#     except:
#         return None

# # ------------------------- GEMINI QA FUNCTION -------------------------
# def generate_answer(question, context):
#     prompt = template.format(question=question, context=context)
#     body = {"contents": [{"parts": [{"text": prompt}]}]}
#     r = requests.post(CHAT_URL, headers={"Content-Type": "application/json"}, json=body)
#     resp = r.json()
#     try:
#         return resp["candidates"][0]["content"]["parts"][0]["text"]
#     except:
#         return "I don't know."

# # ------------------------- DOCUMENT LOADING -------------------------
# def load_pdf(file_path):
#     loader = PDFPlumberLoader(file_path)
#     return loader.load()

# def load_docx(file_path):
#     doc = docx.Document(file_path)
#     full_text = "\n".join([para.text for para in doc.paragraphs])
#     return [{"page_content": full_text}]

# def load_txt(file_path):
#     with open(file_path, "r", encoding="utf-8") as f:
#         text = f.read()
#     return [{"page_content": text}]

# def load_image(file_path):
#     img = Image.open(file_path)
#     text = pytesseract.image_to_string(img)
#     return [{"page_content": text}]

# def split_into_chunks(docs, chunk_size=800, chunk_overlap=200):
#     splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
#     return splitter.split_documents(docs)

# # ------------------------- VECTOR STORE -------------------------
# vector_store = []  # List of {"text": chunk_text, "emb": vector, "source": filename}

# def index_chunks(chunks, filename="unknown"):
#     for ch in chunks:
#         text = ch["page_content"] if isinstance(ch, dict) else ch.page_content
#         emb = get_embedding(text)
#         if emb is not None:
#             vector_store.append({"text": text, "emb": emb, "source": filename})

# def search(query, top_k=3):
#     q_emb = get_embedding(query)
#     if q_emb is None or len(vector_store) == 0:
#         return []

#     scores = []
#     for item in vector_store:
#         sim = cosine_similarity([q_emb], [item["emb"]])[0][0]
#         scores.append((sim, item["text"], item["source"]))

#     scores.sort(reverse=True)
#     return [{"text": t, "source": s} for _, t, s in scores[:top_k]]

# # ------------------------- CHAT MEMORY -------------------------
# if "chat_history" not in st.session_state:
#     st.session_state.chat_history = []

# # ------------------------- STREAMLIT UI -------------------------
# st.set_page_config(page_title="AI Document Assistant", layout="wide")
# st.title("AI Document Assistant — Multi-file / Multi-format")

# # Use existing folder for uploads
# UPLOAD_FOLDER = os.path.join("chat-with-pdf", "pdfs")
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# uploaded_files = st.file_uploader(
#     "Upload PDF, DOCX, TXT, or Image files",
#     type=["pdf","docx","txt","png","jpg","jpeg"],
#     accept_multiple_files=True
# )

# if uploaded_files:
#     for uploaded in uploaded_files:
#         # Save file in 'chat-with-pdf/pdfs' folder
#         path = os.path.join(UPLOAD_FOLDER, uploaded.name)
#         with open(path, "wb") as f:
#             f.write(uploaded.getbuffer())

#         # Load document based on type
#         if path.lower().endswith(".pdf"):
#             docs = load_pdf(path)
#         elif path.lower().endswith(".docx"):
#             docs = load_docx(path)
#         elif path.lower().endswith(".txt"):
#             docs = load_txt(path)
#         elif path.lower().endswith((".png", ".jpg", ".jpeg")):
#             docs = load_image(path)
#         else:
#             st.warning(f"Unsupported file type: {uploaded.name}")
#             continue

#         st.write(f"📄 Processing {uploaded.name}...")
#         chunks = split_into_chunks(docs)
#         index_chunks(chunks, filename=uploaded.name)

#     st.success("✅ All files indexed successfully!")

#     # ------------------------- USER QUERY -------------------------
#     q = st.chat_input("Ask something about your documents")
#     if q:
#         st.chat_message("user").write(q)

#         # Include last 2 Q&A pairs in context
#         past_context = "\n\n".join([f"Q: {c['q']}\nA: {c['a']}" for c in st.session_state.chat_history[-2:]])
#         search_results = search(q)
#         ctx_text = "\n\n".join([f"{r['text']} (Source: {r['source']})" for r in search_results])
#         full_context = past_context + "\n\n" + ctx_text if past_context else ctx_text

#         ans = generate_answer(q, full_context)
#         st.chat_message("assistant").write(ans)

#         # Save to chat history
#         st.session_state.chat_history.append({"q": q, "a": ans})

#     # ------------------------- DOWNLOAD CHAT -------------------------
#     if st.button("📥 Download Chat as PDF"):
#         pdf = FPDF()
#         pdf.add_page()
#         pdf.set_font("Arial", size=12)
#         for item in st.session_state.chat_history:
#             pdf.multi_cell(0, 8, f"Q: {item['q']}")
#             pdf.multi_cell(0, 8, f"A: {item['a']}")
#             pdf.ln()
#         pdf_path = os.path.join(UPLOAD_FOLDER, "chat_history.pdf")
#         pdf.output(pdf_path)
#         with open(pdf_path, "rb") as f:
#             st.download_button("Download PDF", f, file_name="chat_history.pdf", mime="application/pdf")



# University_Assistant.py
import os
import time
import requests
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

import streamlit as st
from html import escape

# For PDFs
from langchain_community.document_loaders.pdf import PDFPlumberLoader
# For DOCX
import docx
# For images (OCR)
from PIL import Image
import pytesseract

# For text splitting
from langchain_text_splitters import RecursiveCharacterTextSplitter

# For exporting chat
from fpdf import FPDF

# ------------------------- CONFIG & THEME -------------------------
st.set_page_config(page_title="Document Assistant", layout="wide", initial_sidebar_state="expanded")

# CSS: Dark, minimal, professional. Supports compact mode via class on body.
CUSTOM_CSS = """
<style>
:root {
  --bg: #0F1117;
  --panel: #0B0C0F;
  --muted: #9aa0a6;
  --text: #E6E6E6;
  --accent: #4F46E5;
}

/* Page container padding */
.block-container {
    padding-top: 1.2rem;
    padding-left: 1.2rem;
    padding-right: 1.2rem;
    background: var(--bg);
}

/* Sidebar */
[data-testid="stSidebar"] {
    background: var(--bg);
    color: var(--text);
    border-right: 1px solid rgba(255,255,255,0.03);
}

/* File card */
.file-card {
    background: linear-gradient(180deg, rgba(255,255,255,0.02), rgba(255,255,255,0.01));
    padding: 12px;
    border-radius: 10px;
    border: 1px solid rgba(255,255,255,0.04);
    margin-bottom: 10px;
    color: var(--text);
    font-size: 14px;
}

/* Chat row and bubbles */
.chat-row {
    display: flex;
    width: 100%;
    margin-bottom: 8px;
}
.user-bubble {
    margin-left: auto;
    background: #20232a;
    color: var(--text);
    padding: 12px 16px;
    border-radius: 14px;
    max-width: 78%;
    box-shadow: 0 1px 6px rgba(0,0,0,0.5);
    font-size: 15px;
    word-wrap: break-word;
}
.assistant-bubble {
    margin-right: auto;
    background: linear-gradient(180deg, #111216, #17181b);
    color: var(--text);
    padding: 12px 16px;
    border-radius: 14px;
    max-width: 78%;
    box-shadow: 0 1px 6px rgba(0,0,0,0.5);
    font-size: 15px;
    word-wrap: break-word;
}

/* Compact mode adjustments */
.compact .file-card { padding: 8px; font-size: 13px; }
.compact .user-bubble, .compact .assistant-bubble { padding: 8px 12px; font-size: 13px; border-radius: 10px; }

/* Small meta text */
.meta {
    color: var(--muted);
    font-size: 12px;
}

/* Buttons style tweaks */
.stButton>button {
    border-radius: 10px;
}

/* Divider */
.hr {
    height: 1px;
    background: rgba(255,255,255,0.03);
    margin: 12px 0;
}

/* Subtle link style for support text */
.support-link {
    color: var(--accent);
    text-decoration: none;
}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# ------------------------- API SETUP -------------------------
API_KEY = os.getenv("GEMINI_API_KEY")
if not API_KEY:
    st.error("GEMINI_API_KEY is not set. Please set it before running the app.")
    st.stop()

EMBED_URL = f"https://generativelanguage.googleapis.com/v1beta/models/text-embedding-004:embedContent?key={API_KEY}"
CHAT_URL  = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={API_KEY}"

# ------------------------- PROMPT TEMPLATE -------------------------
TEMPLATE = """
You are an assistant for question-answering using only the context below.
If the answer is not in the context, say "I don't know."

Question: {question}

Context:
{context}

Answer:
"""

# ------------------------- CORE FUNCTIONS -------------------------
def get_embedding(text):
    headers = {"Content-Type": "application/json"}
    body = {"model": "models/text-embedding-004", "content": {"parts": [{"text": text}]}}
    try:
        r = requests.post(EMBED_URL, headers=headers, json=body, timeout=30)
        r.raise_for_status()
        emb = r.json().get("embedding", {}).get("values")
        if emb:
            return np.array(emb)
    except Exception:
        # Silent on-screen: show neutral warning once
        st.warning("Embedding service unavailable. Try again later.")
    return None

def generate_answer(question, context):
    prompt = TEMPLATE.format(question=question, context=context)
    body = {"contents": [{"parts": [{"text": prompt}]}]}
    try:
        r = requests.post(CHAT_URL, headers={"Content-Type": "application/json"}, json=body, timeout=60)
        r.raise_for_status()
        resp = r.json()
        return resp["candidates"][0]["content"]["parts"][0]["text"]
    except Exception:
        st.warning("Answer generation failed. Try again.")
        return "I don't know."

# ------------------------- DOCUMENT LOADERS -------------------------
def load_pdf(file_path):
    loader = PDFPlumberLoader(file_path)
    return loader.load()

def load_docx(file_path):
    doc = docx.Document(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    return [{"page_content": full_text}]

def load_txt(file_path):
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    return [{"page_content": text}]

def load_image(file_path):
    img = Image.open(file_path)
    text = pytesseract.image_to_string(img)
    return [{"page_content": text}]

def split_into_chunks(docs, chunk_size=800, chunk_overlap=200):
    splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    return splitter.split_documents(docs)

# ------------------------- INDEXING -------------------------
def index_chunks(chunks, filename, progress_callback=None):
    """Convert chunks to embeddings and store in vector store."""
    added = 0
    total = len(chunks)

    for i, ch in enumerate(chunks, start=1):
        text = ch.page_content if hasattr(ch, "page_content") else ch.get("page_content", "")

        emb = get_embedding(text)
        if emb is None:
            continue

        st.session_state.vector_store.append({
            "text": text,
            "emb": emb,
            "source": filename
        })

        added += 1
        if progress_callback:
            progress_callback(i, total)

    return added

# ------------------------- SEARCH -------------------------
def search(query, top_k=5):
    """Return top K most similar chunks from vector store."""
    if len(st.session_state.vector_store) == 0:
        return []

    q_emb = get_embedding(query)
    if q_emb is None:
        return []

    sims = []
    for item in st.session_state.vector_store:
        sim = cosine_similarity([q_emb], [item["emb"]])[0][0]
        sims.append((sim, item))

    sims.sort(key=lambda x: x[0], reverse=True)

    return [item for _, item in sims[:top_k]]

# ------------------------- SESSION STATE -------------------------
if "vector_store" not in st.session_state:
    st.session_state.vector_store = []   # {"text","emb","source"}
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []   # {"q","a"}
if "last_index_time" not in st.session_state:
    st.session_state.last_index_time = None
if "compact_mode" not in st.session_state:
    st.session_state.compact_mode = False

# ------------------------- UPLOAD FOLDER -------------------------
UPLOAD_FOLDER = os.path.join("chat-with-pdf", "pdfs")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ------------------------- SIDEBAR (non-destructive, helpful info) -------------------------
with st.sidebar:
    st.markdown("## Document Assistant")
    st.markdown("<div class='meta'>Upload documents and ask questions</div>", unsafe_allow_html=True)
    st.divider()

    
    st.markdown("**How to use**")
    st.markdown("1. Go to Files → upload supported documents.\n2. Wait for the app to finish processing.\n3. Go to Chat → ask questions about uploaded content.\n4. Download chat if needed from Download tab.")
    st.divider()
    st.markdown("**Supported formats**")
    st.markdown("- PDF, DOCX, TXT, PNG, JPG")
    st.divider()
    

# Toggle compact class on body by injecting a small script (works visually)
if st.session_state.compact_mode:
    st.markdown("<script>document.querySelector('body').classList.add('compact');</script>", unsafe_allow_html=True)
else:
    st.markdown("<script>document.querySelector('body').classList.remove('compact');</script>", unsafe_allow_html=True)

# ------------------------- MAIN UI -------------------------
st.markdown("<h2 style='color:#E6E6E6; margin-bottom:4px'>Document Assistant</h2>", unsafe_allow_html=True)
st.markdown("<div class='meta'>Upload documents and ask questions</div>", unsafe_allow_html=True)
st.markdown("<div class='hr'></div>", unsafe_allow_html=True)

tab_files, tab_chat, tab_download, tab_about = st.tabs(["Files", "Chat", "Download", "About"])

# ---------- FILES TAB ----------
with tab_files:
    st.subheader("Upload documents")
    st.markdown("Supported formats: PDF, DOCX, TXT, PNG, JPG. Files are saved locally for processing (not listed here).")

    uploaded_files = st.file_uploader(
        "Select files",
        type=["pdf", "docx", "txt", "png", "jpg", "jpeg"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        overall_progress = st.progress(0)
        total_files = len(uploaded_files)
        indexed_this_run = 0
        file_index = 0

        for uploaded in uploaded_files:
            file_index += 1
            fname = uploaded.name
            # Save locally (keeps behavior consistent). Files are not displayed anywhere.
            path = os.path.join(UPLOAD_FOLDER, fname)
            with open(path, "wb") as f:
                f.write(uploaded.getbuffer())

            # Read file contents
            try:
                if path.lower().endswith(".pdf"):
                    docs = load_pdf(path)
                elif path.lower().endswith(".docx"):
                    docs = load_docx(path)
                elif path.lower().endswith(".txt"):
                    docs = load_txt(path)
                elif path.lower().endswith((".png", ".jpg", ".jpeg")):
                    docs = load_image(path)
                else:
                    st.warning("Unsupported file type.")
                    continue
            except Exception:
                st.warning(f"Could not read {fname}.")
                continue

            # Neutral file card
            st.markdown(
                f"""
                <div class="file-card">
                    📄 <strong>{escape(fname)}</strong><br>
                    <div class="meta">Reading file…</div>
                </div>
                """,
                unsafe_allow_html=True
            )

            # Chunk & index
            chunks = split_into_chunks(docs)
            n_chunks = len(chunks)

            def progress_cb(i, total):
                pct = min(100, int((file_index - 1) / total_files * 100 + (i / total_files / total) * 100))
                overall_progress.progress(pct)

            added = 0
            try:
                added = index_chunks(chunks, filename=fname, progress_callback=progress_cb)
            except Exception:
                added = 0

            indexed_this_run += added
            overall_progress.progress(int(file_index / total_files * 100))

            # Succinct success card
            st.markdown(
                f"""
                <div class="file-card">
                    📄 <strong>{escape(fname)}</strong><br>
                    <span class="meta">Indexed successfully</span>
                </div>
                """,
                unsafe_allow_html=True
            )

        # Update last indexed time and show subtle completion message
        st.session_state.last_index_time = time.strftime("%Y-%m-%d %H:%M:%S")
        st.success("Files processed.")

# ---------- CHAT TAB ----------
with tab_chat:
    st.subheader("Ask questions")

    if len(st.session_state.vector_store) == 0:
        st.info("No documents indexed yet. Upload files in the Files tab first.")
    else:
        # show last conversation (recent)
        st.markdown("<div style='margin-bottom:8px' class='meta'>Recent conversation</div>", unsafe_allow_html=True)
        for msg in st.session_state.chat_history[-6:]:
            q = msg.get("q")
            a = msg.get("a")
            if q:
                st.markdown(f"<div class='chat-row'><div class='user-bubble'>{escape(q)}</div></div>", unsafe_allow_html=True)
            if a:
                st.markdown(f"<div class='chat-row'><div class='assistant-bubble'>{escape(a)}</div></div>", unsafe_allow_html=True)

        query = st.text_input("Ask about your documents", value="", key="chat_input")
        submit = st.button("Ask")

        if submit and query:
            # render user bubble instantly
            st.markdown(f"<div class='chat-row'><div class='user-bubble'>{escape(query)}</div></div>", unsafe_allow_html=True)

            # include last 2 Q&A pairs in context for continuity
            past_context = "\n\n".join([f"Q: {c['q']}\nA: {c['a']}" for c in st.session_state.chat_history[-2:]])
            search_results = search(query)
            if not search_results:
                st.markdown("<div class='assistant-bubble'>I don't know.</div>", unsafe_allow_html=True)
                st.session_state.chat_history.append({"q": query, "a": "I don't know."})
            else:
                ctx_text = "\n\n".join([r["text"] for r in search_results])
                full_context = (past_context + "\n\n" + ctx_text) if past_context else ctx_text
                with st.spinner("Working..."):
                    ans = generate_answer(query, full_context)
                st.markdown(f"<div class='chat-row'><div class='assistant-bubble'>{escape(ans)}</div></div>", unsafe_allow_html=True)
                st.session_state.chat_history.append({"q": query, "a": ans})

# ---------- DOWNLOAD TAB ----------
with tab_download:
    st.subheader("Export")
    st.markdown("Download the chat history as a PDF (local).")

    if st.button("Download chat as PDF"):
        if len(st.session_state.chat_history) == 0:
            st.info("No chat history available.")
        else:
            pdf = FPDF()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            for item in st.session_state.chat_history:
                pdf.multi_cell(0, 8, f"Q: {item['q']}")
                pdf.multi_cell(0, 8, f"A: {item['a']}")
                pdf.ln()
            pdf_path = os.path.join(UPLOAD_FOLDER, "chat_history.pdf")
            pdf.output(pdf_path)
            with open(pdf_path, "rb") as f:
                st.download_button("Download PDF", f, file_name="chat_history.pdf", mime="application/pdf")

# ---------- ABOUT TAB ----------
with tab_about:
    st.subheader("About")
    st.markdown(
        """
        - Upload PDF, DOCX, TXT or images.  
        - The app extracts text and answers questions from what you upload.
        """
    )
    st.markdown("Notes")
    st.markdown("- Ensure `GEMINI_API_KEY` is set in the environment before running.")